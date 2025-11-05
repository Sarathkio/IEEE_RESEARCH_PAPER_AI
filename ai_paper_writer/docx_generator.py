from docx import Document
from io import BytesIO

def create_docx_bytes(sections, bib_entries):
    """
    Generates a DOCX document from the given sections and bibliography entries.

    Args:
        sections (dict): Dictionary containing section titles and content.
        bib_entries (list): List of bibliography reference strings.

    Returns:
        bytes: DOCX file in bytes format for downloading or saving.
    """
    
    # Create a new Word document
    doc = Document()
    
    # Add the main title
    doc.add_heading("Generated Research Paper", 0)

    # Add each section with heading and content
    for sec, content in sections.items():
        doc.add_heading(sec.capitalize(), level=1)   # Section heading
        doc.add_paragraph(content)                   # Section content

    # Add the References section
    doc.add_heading("References", level=1)
    for ref in bib_entries:
        doc.add_paragraph(ref)                       # Each reference as a paragraph

    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Move cursor to start of the buffer

    return buffer.read()  # Return byte content of the document
