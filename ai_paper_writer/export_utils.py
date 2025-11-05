# export_utils.py
from docx import Document
from fpdf import FPDF

def export_to_docx(sections: dict) -> bytes:
    doc = Document()
    for heading, content in sections.items():
        if heading == "Full Content":
            continue
        doc.add_heading(heading, level=1)
        for para in content.split('\n\n'):
            doc.add_paragraph(para.strip())
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def export_to_pdf(sections: dict) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for heading, content in sections.items():
        if heading == "Full Content":
            continue
        pdf.set_font(style="B", size=14)
        pdf.cell(200, 10, txt=heading, ln=True)
        pdf.set_font("Arial", size=12)
        for para in content.split('\n\n'):
            pdf.multi_cell(0, 10, para.strip())
    buffer = io.BytesIO()
    pdf.output(buffer)
    return buffer.getvalue()

def export_to_markdown(sections: dict) -> str:
    md = ""
    for heading, content in sections.items():
        if heading == "Full Content":
            continue
        md += f"# {heading}\n\n{content}\n\n"
    return md
